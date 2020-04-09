#!/usr/bin/env python3
# -*- coding=utf-8 -*-

# @Time    : 2018-11-02
# @Author  : J.sky
# @Mail    : bosichong@qq.com
# @Site    : www.17python.com
# @Title   : 基于Python开发的小学生口算题生成器
# @Url     : http://www.17python.com/blog/29
# @Details : Python实现小学生加减乘除速算考试题卷。
# @Other   : OS X 10.11.6
#            Python 3.6.1
#            PyCharm


'''
孩子上小学一年级了，加减乘除的口算就要开始练习了，估计老题肯定会让家长出题，所以提前准备一下.

利用Python开发了一套自动生成小学生口算题的小应用。而且今天是程序员节，撸200行代码庆祝一下。：）

程序核心功能：

    1.根据条件生成相关的口算题.

    2.保存为.docx用来打印.


开心Python Django 学习交流q群：217840699


Author  : J.sky
Mail    : bosichong@qq.com


'''

from docx import Document  # 引入docx类生成docx文档
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.shared import RGBColor, Cm
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

__version__ = "1.0.0"


class PrintPreview:
    '''本类负责生成完整的口算题文档使之适合打印机打印。可以生成多套题，生成数可以控。

    - @p_list   list
    需要打印口算题库，至少包含一套口算题

    - @p_title   list
    页面标题，这个标题的生成依据程序题型的选择和数字的范围选择而生成，例如：选择了0-20，加减法，进退位
    则自动生成标题为：0到20加减法进退位混合口算题，list中包含了多套题的页面标题名称

    - @p_column  int
    打印页排版口算题的列数

    '''

    p_list = None
    p_title = None
    p_subtitle = None
    p_column = None
    p_title_size = None
    p_subtitle_size = None
    p_content_siae = None

    def __init__(self, l, tit, subtitle, col=2, tsize=20, subsize=16, csize=18):
        '''
        :param l: list 需要打印的口算题列表
        :param tit: list 口算页标题
        :param subtitle str 小标题
        :param col: int 列数
        :param tsize: int 标题字号
        :param csize: int 口算题字号
        '''
        self.p_list = l
        self.p_title = tit
        self.p_subtitle = subtitle
        self.p_column = col
        self.p_title_size = tsize
        self.p_subtitle_size = subsize
        self.p_content_siae = csize

    def create_psmdocx(self, expressionList, title, docxname):
        '''
        :param expressionList list 一组题库
        :param title str 页面标题
        :param docxname  str 题库保存文件名
        :return: none
        '''
        if (title == ''):
            page_title = '算术题'
        else:
            page_title = title
        p_docx = Document()  # 创建一个docx文档
        # 自定义正文格式
        p_docx.styles['Normal'].font.name = u'Arial'  # 可换成word里面任意字体
        p_docx.styles['Normal'].paragraph_format.space_before = Pt(5)
        p_docx.styles['Normal'].paragraph_format.space_after = Pt(5)
        p_docx.styles['Normal'].font.size = Pt(self.p_content_siae)
        # 自定义标题格式
        p_docx.styles['Heading 1'].paragraph_format.space_before = Pt(12)
        p_docx.styles['Heading 1'].paragraph_format.space_after = Pt(12)
        p_docx.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落文字居中设置

        # 判断需要用到的行数
        if (len(expressionList) % self.p_column == 0):
            rs = len(expressionList) // self.p_column
        else:
            rs = len(expressionList) // self.p_column + 1

        # print(rs)

        # 将口算题添加到docx表格中
        # 每页10行
        tableRows = 10
        for i in range(rs):
            if (i % tableRows == 0):
                if i > 0:
                    # 添加分页符
                    p_docx.add_page_break()
                # 添加页头内容
                self.addPageHeader(p_docx, page_title)
                # 添加算式到表格
                table = p_docx.add_table(rows=tableRows, cols=self.p_column)
                # 自定义行高
                for row in table.rows:
                    row.height = Cm(1.8)
                table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table.style.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
                table.style.font.size = Pt(self.p_content_siae)  # 字体大小设置，和word里面的字号相对应
                table.rows.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                k = 0  # 计数器

            # for i in range(rs):
            if i%tableRows >= 0:
                row_cells = table.rows[i%tableRows].cells
                for j in range(self.p_column):
                    columnIndex = 2 * i + j
                    if (columnIndex > len(expressionList) - 1):
                        print('第{}行、第{}列，超出算式列表总数{}'.format(i, columnIndex, len(expressionList)))
                        break
                    else:
                        row_cells[j].text = expressionList[columnIndex]
                        k = k + 1

        p_docx.save('{}.docx'.format(docxname))  # 输出docx

    def addPageHeader(self, p_docx, page_title):
        # 新建标题并应用格式
        paragraph_title = p_docx.add_paragraph()
        paragraph_title.style = p_docx.styles['Heading 1']
        title = paragraph_title.add_run(page_title)
        title.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        title.font.size = Pt(self.p_title_size)  # 字体大小设置，和word里面的字号相对应
        title.font.name = u'楷体'
        title = self.setZhFont(title, u'楷体')
        # 新建子标题，设置字体大小
        subTitle = p_docx.add_paragraph()
        subTitle.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落文字居中设置
        srun = subTitle.add_run(self.p_subtitle)
        srun.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        srun.font.size = Pt(self.p_subtitle_size)  # 字体大小设置，和word里面的字号相对应
        srun.font.name = u'Arial'
        srun = self.setZhFont(srun, u'楷体')
        # 添加空行
        p_docx.add_paragraph()

    def setZhFont(self, run, zhFontName):
        """
        设置字体，含西文（数字）字体、中文字体
        :param run: 待设置的行
        :param fontName: 西文（数字）字体
        :param zhFontName: 中文字体
        :return:
        """
        run._element.rPr.rFonts.set(qn('w:eastAsia'), zhFontName)
        return run


    def produce(self):
        '''
        生成.docx文档
        :return:
        '''
        k = 1
        for l, t in zip(self.p_list, self.p_title):
            self.create_psmdocx(l, t, t + str(k))
            k = k + 1



if __name__ == '__main__':
    l = [['1-17=', '3-4=', '13-6=', '15-5=', '2-4=', '15-9=', '12-13=', '15-12=', '14-16=', '4-11=', '18-16=', '12-14=',
          ],
         ['1-17=', '3-4=', '13-6=', '15-5=', '2-4=', '15-9=', '12-13=', '15-12=', '14-16=', '4-11=', '18-16=', '12-14=',
          '14-7=', '7-17=', '16-19=',  ]]
    t = ['小学生口算题', '小学生口算题']
    pp = PrintPreview(l, t,"姓名：__________ 日期：____月____日 时间：________ 对题：____道" ,4)
    pp.produce()
